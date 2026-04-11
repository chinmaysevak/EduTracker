import collections
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_font(run, size=12, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold

def add_heading(doc, text, level, size, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    if level == 0:
        p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    apply_font(run, size=size, bold=True)
    return p

def add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    apply_font(run, size=12, bold=bold)
    return p

def main():
    doc = Document()

    # Set paper size to A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    # Set margins
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Change Normal style font and line spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Cover Page
    add_heading(doc, 'A MINI PROJECT REPORT', level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading(doc, 'ON', level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading(doc, '"EduTrack – Student Assistant"', level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.add_paragraph() # space
    
    p = add_paragraph(doc, 'Submitted in the partial fulfillment of the requirements for', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    p = add_paragraph(doc, 'The degree of', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    
    doc.add_paragraph()
    add_heading(doc, 'BACHELOR of COMPUTER APPLICATION', level=1, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.add_paragraph()
    add_paragraph(doc, 'By', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, '1) ______________________________', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '2) ______________________________', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '3) ______________________________', align=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.add_paragraph()
    add_paragraph(doc, 'UNDER THE GUIDANCE OF', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, 'Dr. Anjali Dadhich', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, 'Department of Management Studies', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, 'Bharati Vidyapeeth Deemed to be University', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, '2025-26', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    doc.add_page_break()

    # Vision & Mission
    add_heading(doc, 'Vision', level=1, size=16)
    add_paragraph(doc, 'Preparing the Students to cope with the rigor of Post Graduate Programmes in global and creating high caliber solution architects for software development, who will also be sensitive to societal concerns.')
    doc.add_paragraph()
    add_heading(doc, 'Mission', level=1, size=16)
    add_paragraph(doc, '- We aim to drive transformation, technology and innovation through problem solving approach and research development.')
    add_paragraph(doc, '- We aim to provide students with the IT tools to become productive and lifelong learner.')

    doc.add_page_break()

    # Certificate
    add_heading(doc, 'CERTIFICATE', level=1, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, 'This is to certify that the requirements for the project report entitled "EduTrack – Student Assistant" have been successfully completed by the following students: ')
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Roll Numbers'
    hdr_cells[1].text = 'Name'
    for cell in hdr_cells:
        apply_font(cell.paragraphs[0].runs[0], bold=True)
    
    add_paragraph(doc, 'In partial fulfillment of Sem – IV, Bachelor of Computer Application of "Bharati Vidyapeeth Deemed to be University, Department of Management Studies", Kharghar during the academic year 2025-26.')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    table2 = doc.add_table(rows=2, cols=2)
    cell00 = table2.cell(0, 0)
    p = cell00.paragraphs[0]
    r = p.add_run('Project Co-ordinator')
    apply_font(r, bold=True)
    
    cell01 = table2.cell(0, 1)
    p = cell01.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('External Examiner')
    apply_font(r, bold=True)
    
    cell10 = table2.cell(1, 0)
    cell10.text = 'Dr. Anjali Dadhich'
    apply_font(cell10.paragraphs[0].runs[0])
    
    doc.add_paragraph()
    add_paragraph(doc, 'Director', bold=True)
    add_paragraph(doc, 'Dr. Premashish Roy')

    doc.add_page_break()

    # PEOs
    add_heading(doc, "PROGRAM EDUCATIONAL OBJECTIVE's", level=1, size=16)
    add_paragraph(doc, 'I. To prepare the youth to take up positions as system analysts, system engineers, software engineers and programmers.')
    add_paragraph(doc, "II. To aim at developing 'systems thinking' 'abstract thinking', 'skills to analyze and synthesize', and 'skills to apply knowledge', through 'extensive problem solving sessions', 'hands on practice under various hardware/software environments' and 'projects developed'.")
    add_paragraph(doc, "III. To prepare students with 'social interaction skills', 'communication skills', 'life skills', 'entrepreneurial skills'.")

    doc.add_page_break()

    # POs
    add_heading(doc, 'PROGRAM OUTCOMES', level=1, size=16)
    pos = [
        "PO1. Computational Knowledge: Understand and apply mathematical foundation, computing and domain knowledge for the conceptualization of computing models from defined problems.",
        "PO2. Problem Analysis: Ability to identify, critically analyze and formulate complex computing problems using fundamentals of computer science and application domains.",
        "PO3. Design / Development of Solutions: Ability to transform complex business scenarios and contemporary issues into problems, investigate, understand and propose integrated solutions using emerging technologies.",
        "PO4. Conduct Investigations of Complex Computing Problems: Ability to devise and conduct experiments, interpret data and provide well informed conclusions.",
        "PO5. Modern Tool Usage: Ability to select modern computing tools, skills and techniques necessary for innovative software solutions.",
        "PO6. Professional Ethics: Ability to apply and commit professional ethics and cyber regulations in a global economic environment.",
        "PO7. Life-long Learning: Recognize the need for and develop the ability to engage in continuous learning as a Computing professional.",
        "PO8. Project Management: Ability to understand management and computing principles with computing knowledge to manage projects in multidisciplinary environments.",
        "PO9. Communication Efficacy: Communicate effectively with the computing community as well as society by being able to comprehend effective documentations and presentations.",
        "PO10. Societal & Environmental Concern: Ability to recognize economical, environmental, social, health, legal, ethical issues involved in the use of computer technology and other consequential responsibilities relevant to professional practice.",
        "PO11. Individual & Team Work: Ability to work as a member or leader in diverse teams in multidisciplinary environment.",
        "PO12. Innovation and Entrepreneurship: Identify opportunities, entrepreneurship vision and use of innovative ideas to create value and wealth for the betterment of the individual and society."
    ]
    for po in pos:
        add_paragraph(doc, po)

    doc.add_page_break()

    # PSOs
    add_heading(doc, 'PROGRAMME SPECIFIC OUTCOME', level=1, size=16)
    add_paragraph(doc, 'After the completion of the course, a student is able to:')
    psos = [
        "PSO1: Ability to learn the various programming languages with database concepts along with development environment.",
        "PSO2: Ability to apply theoretical and practical knowledge to solve business problems through data communication technology concepts.",
        "PSO3: Flourish the innovation and research attitude to develop IT artifact.",
        "PSO4: Foster analytical and critical thinking abilities for efficient programming.",
        "PSO5: Demonstrate and apply the programming knowledge to develop effective software solution.",
        "PSO6: Enrich the knowledge in the areas of Advanced technologies and business practices.",
        "PSO7: Maintain the personality with environmental and social concerns."
    ]
    for pso in psos:
        add_paragraph(doc, pso)

    doc.add_page_break()

    # Declaration
    add_heading(doc, 'DECLARATION', level=1, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "I declare that this written submission represents my ideas in my own words and where others' ideas or words have been included. I have adequately cited and referenced the original sources. I also declare that I have adhered to all principles of academic honesty and integrity and have not misrepresented or fabricated or falsified any idea/data/fact/source in my submission. I understand that any violation of the above will be cause for disciplinary action by the Institute and can also evoke penal action from the sources which have thus not been properly cited or from whom proper permission has not been taken when needed.")
    doc.add_paragraph()
    add_paragraph(doc, '1) ______________________________')
    add_paragraph(doc, '2) ______________________________')
    add_paragraph(doc, '3) ______________________________')
    doc.add_paragraph()
    add_paragraph(doc, 'Date: ', bold=True)

    doc.add_page_break()

    # Acknowledgement
    add_heading(doc, 'ACKNOWLEDGEMENT', level=1, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "After the completion of this work, words are not enough to express feelings about all those who helped us to reach our goal.")
    add_paragraph(doc, "It's a great pleasure and moment of immense satisfaction for us to express our profound gratitude to our Project Guide, Prof. Dr. Anjali Dadhich, whose constant encouragement enabled us to work enthusiastically. Her perpetual motivation, patience and excellent expertise in discussion during the progress of the project work have benefited us to an extent, which is beyond expression.")
    add_paragraph(doc, "We would also like to give our sincere thanks to Dr. Mona Sinha, Head of Department, and Dr. Anjali Dadhich, Project Co-ordinator from Department of Management Studies, Kharghar, for their guidance, encouragement, and support during the project.")
    add_paragraph(doc, "We are thankful to Dr. Premashish Roy, Director, for providing an outstanding academic environment, also for providing the adequate facilities.")
    add_paragraph(doc, "Last but not the least we would also like to thank all the staff of Bharati Vidyapeeth Deemed to be University, Department of Management Studies for their valuable guidance with their interest and valuable suggestions which brightened us.")
    doc.add_paragraph()
    add_paragraph(doc, '1) ______________________________')
    add_paragraph(doc, '2) ______________________________')

    doc.add_page_break()

    # Abstract
    add_heading(doc, 'Abstract', level=1, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "EduTrack – Student Assistant is a comprehensive, full-stack web application designed to empower students in managing their academic life efficiently. Built using the MERN stack (MongoDB, Express.js, React, Node.js) with TypeScript on the frontend, EduTrack provides an all-in-one platform for attendance tracking, study planning, task management, focus sessions, progress monitoring, resource management, and AI-powered academic guidance.")
    add_paragraph(doc, "The application addresses the critical challenge faced by modern students: the lack of a unified, intelligent platform to track and optimize their academic activities. Traditional methods of attendance tracking, study planning, and progress monitoring are fragmented and inefficient. EduTrack consolidates all these functionalities into a single, beautifully designed Progressive Web App (PWA) that works seamlessly across desktop and mobile devices.")
    add_paragraph(doc, "Key features include a real-time dashboard with smart widgets, subject-wise attendance tracking with percentage calculations, a study planner with task management and exam countdown, an AI-powered Smart Advisor using Google's Gemini API, a Focus Mode with Pomodoro timer and ambient sounds, syllabus progress tracking, resource management (notes, links, YouTube videos, files), gamification with XP points and achievement badges, Google OAuth integration for seamless authentication, and a robust import/export system for data portability.")
    add_paragraph(doc, "The backend is powered by Node.js with Express.js, using MongoDB Atlas as the cloud database with Mongoose ODM for data modeling. Authentication is handled through JWT tokens with support for both email/password and Google Sign-In. The frontend leverages React 19 with TypeScript, Tailwind CSS, Radix UI primitives, Framer Motion animations, and Recharts for data visualization.")
    doc.add_paragraph()
    p = add_paragraph(doc, "Keywords: ", bold=True)
    p.add_run("Student Management System, MERN Stack, Attendance Tracker, Study Planner, Progressive Web App, AI Academic Assistant, React, MongoDB, TypeScript.")

    doc.add_page_break()

    # Table of Contents
    add_heading(doc, 'Table of Contents', level=1, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    table_toc = doc.add_table(rows=15, cols=2)
    tdata = [
        ("List of Figures", "1"),
        ("1. Introduction", "2"),
        ("    1.1 General", "2"),
        ("    1.2 Objective and Problem Statement", "3"),
        ("2. Methodology", "6"),
        ("    2.1 Algorithmic Details", "6"),
        ("    2.2 Hardware and Software Requirements", "12"),
        ("    2.3 Design Details", "13"),
        ("3. Implementation and Results", "15"),
        ("    3.1 Implementation", "15"),
        ("    3.2 Results", "16"),
        ("4. Conclusion and Future Scope", "18"),
        ("    4.1 Conclusion", "18"),
        ("    4.2 Future Scope", "19"),
        ("5. References", "19")
    ]
    for i, (col1, col2) in enumerate(tdata):
        table_toc.cell(i, 0).text = col1
        table_toc.cell(i, 1).text = col2
        apply_font(table_toc.cell(i, 0).paragraphs[0].runs[0])
        apply_font(table_toc.cell(i, 1).paragraphs[0].runs[0])

    doc.add_page_break()

    # List of figures
    add_heading(doc, 'List of Figures', level=1, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    table_lof = doc.add_table(rows=17, cols=3)
    table_lof.style = 'Table Grid'
    lof_data = [
        ("Figure No.", "Name", "Page No."),
        ("Fig 1.1", "System Architecture Diagram", "3"),
        ("Fig 2.1", "User Authentication Flowchart", "7"),
        ("Fig 2.2", "Attendance Tracking Algorithm", "8"),
        ("Fig 2.3", "Focus Session State Machine", "9"),
        ("Fig 2.4", "AI Smart Advisor Data Flow", "10"),
        ("Fig 2.5", "Study Planner Task Lifecycle", "11"),
        ("Fig 2.6", "Entity-Relationship Diagram", "14"),
        ("Fig 2.7", "Data Flow Diagram (DFD Level 0)", "14"),
        ("Fig 2.8", "Data Flow Diagram (DFD Level 1)", "15"),
        ("Fig 3.1", "Login Page Screenshot", "16"),
        ("Fig 3.2", "Dashboard Screenshot", "16"),
        ("Fig 3.3", "Attendance Tracker Screenshot", "17"),
        ("Fig 3.4", "Study Planner Screenshot", "17"),
        ("Fig 3.5", "Focus Mode Screenshot", "18"),
        ("Fig 3.6", "AI Smart Advisor Screenshot", "18"),
        ("Fig 3.7", "Settings Page Screenshot", "18")
    ]
    for i, (col1, col2, col3) in enumerate(lof_data):
        table_lof.cell(i, 0).text = col1
        table_lof.cell(i, 1).text = col2
        table_lof.cell(i, 2).text = col3
        for cell in table_lof.rows[i].cells:
            if cell.paragraphs[0].runs:
                apply_font(cell.paragraphs[0].runs[0], bold=(i==0))

    doc.add_page_break()

    # Chapters Placeholder Setup
    def add_chapter(title, subtitle):
        add_heading(doc, f'CHAPTER', level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER) # usually chapter number goes here
        add_heading(doc, title, level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
        if subtitle:
            add_heading(doc, subtitle, level=1, size=16)

    # Note: Full contents for chapters are long, I will append them using a generic read
    # from a text file that I will also create, to avoid massive script file size, or I can just chunk it.
    
    # Rather than parsing plain text in python-docx, let's just write everything here 
    # to be safe and deterministic.
    
    add_heading(doc, 'CHAPTER 1', level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading(doc, 'INTRODUCTION', level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading(doc, '1.1 GENERAL', level=1, size=16)
    
    add_paragraph(doc, "In today's digital age, students face an ever-increasing challenge of managing their academic activities efficiently. From tracking daily attendance across multiple subjects to planning study sessions, managing assignments, and monitoring overall progress—the modern student juggles a multitude of responsibilities that are often tracked using fragmented, disconnected tools such as spreadsheets, paper notebooks, and separate mobile apps.")
    add_paragraph(doc, "EduTrack – Student Assistant is a modern, full-stack web application that addresses this gap by providing a unified, intelligent, and aesthetically pleasing platform for comprehensive student academic management. The application is built using the MERN stack (MongoDB, Express.js, React, Node.js), which represents one of the most popular and powerful technology stacks for building modern web applications.")
    add_paragraph(doc, "The application leverages the latest web technologies including React 19 with TypeScript for type-safe frontend development, Tailwind CSS for responsive styling, Radix UI for accessible component primitives, and Framer Motion for fluid animations. On the backend, Express.js serves as the REST API framework, MongoDB Atlas provides cloud-based database storage, and Mongoose serves as the Object Data Modeling (ODM) library.")
    add_paragraph(doc, "EduTrack operates as a Progressive Web App (PWA), meaning it can be installed on mobile devices and desktops, works offline with service workers, and provides a native app-like experience. The application supports Google OAuth 2.0 authentication alongside traditional email/password login, secured with JWT (JSON Web Tokens) and bcrypt password hashing.")
    add_paragraph(doc, "A standout feature of EduTrack is its AI-powered Smart Advisor, which integrates with Google's Gemini API to provide intelligent academic guidance, study recommendations, and personalized tips based on the student's performance data.")
    add_paragraph(doc, "The system architecture follows a clear client-server model with the React frontend deployed on Vercel/Netlify and the Express.js backend deployed separately, communicating through RESTful API endpoints. Data is persisted in MongoDB Atlas with 15 distinct database collections covering all aspects of student academic life.")
    
    doc.add_paragraph()
    
    add_heading(doc, '1.2 OBJECTIVE AND PROBLEM STATEMENT', level=1, size=16)
    add_heading(doc, 'Problem Statement', level=2, size=14)
    add_paragraph(doc, "Students across educational institutions face significant challenges in managing their academic life effectively:")
    add_paragraph(doc, "1. Fragmented Tracking Tools: Attendance is tracked manually or in separate spreadsheets; tasks are managed in different apps; study materials are scattered across multiple platforms. There is no unified system that brings all academic management under one roof.")
    add_paragraph(doc, "2. Lack of Intelligent Insights: Traditional tools provide no smart analysis. Students cannot easily determine how many classes they can miss to maintain a required attendance percentage, or which subjects need more focus based on their progress data.")
    add_paragraph(doc, "3. No Focused Study Environment: Students struggle with procrastination and lack tools that provide structured study sessions with timers, ambient sounds, and distraction-free environments.")
    add_paragraph(doc, "4. Inaccessible Across Devices: Many existing solutions are platform-specific (only Android/iOS) or require installation of heavy software. Students need a solution that works seamlessly across all devices—laptops, tablets, and phones.")
    add_paragraph(doc, "5. No AI-Powered Guidance: Students lack an intelligent assistant that can analyze their academic data and provide personalized study recommendations.")
    add_paragraph(doc, "6. Data Portability: Existing solutions lock students into their ecosystems with no easy way to export or back up their academic data.")
    
    add_heading(doc, 'Objectives', level=2, size=14)
    add_paragraph(doc, "The primary objectives of the EduTrack project are:")
    add_paragraph(doc, "1. Develop a Unified Academic Management Platform: Create a single web application that handles attendance tracking, study planning, task management, resource organization, progress monitoring, and academic analytics.")
    add_paragraph(doc, "2. Implement Intelligent Attendance Analysis: Build algorithms that calculate subject-wise attendance percentages, determine the number of classes a student can miss to maintain a target percentage, and provide smart alerts for low attendance.")
    add_paragraph(doc, "3. Create a Focus-Enhancing Study Environment: Develop a Focus Mode with Pomodoro technique timer, customizable study/break durations, ambient sounds (rain, lo-fi music, café ambiance, fire, forest), and distraction-free interface.")
    add_paragraph(doc, "4. Integrate AI-Powered Academic Advisor: Utilize Google's Gemini AI API to provide personalized study advice, tips, and recommendations based on the student's academic data including attendance records, task completion rates, and progress metrics.")
    add_paragraph(doc, "5. Build a Progressive Web App (PWA): Ensure the application is installable, works offline, provides push notifications for class reminders and attendance alerts, and delivers a native app-like experience.")
    add_paragraph(doc, "6. Implement Secure Authentication: Provide multi-method authentication using JWT-based email/password login, Google OAuth 2.0 Sign-In, OTP-based email verification, and password reset functionality.")
    add_paragraph(doc, "7. Enable Data Portability: Allow complete import/export of all user data in JSON format for backup, restore, and portability.")
    add_paragraph(doc, "8. Implement Gamification: Introduce XP points, levels, streaks, and achievement badges to motivate students to maintain consistent study habits.")
    add_paragraph(doc, "9. Design a Premium, Responsive UI: Create a visually stunning interface with glassmorphism effects, smooth animations, dark/light theme support, and full responsiveness across all screen sizes.")
    add_paragraph(doc, "10. Deploy on Cloud Infrastructure: Deploy the frontend on Vercel/Netlify and the backend on a cloud server with MongoDB Atlas for reliable, globally accessible service.")

    doc.add_page_break()

    add_heading(doc, 'CHAPTER 2', level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading(doc, 'METHODOLOGY', level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_heading(doc, '2.1 ALGORITHMIC DETAILS', level=1, size=16)
    add_heading(doc, 'Algorithm 1: User Authentication Flow', level=2, size=14)
    add_paragraph(doc, "This algorithm outlines the flow for both Email/Password and Google OAuth authentication methods.")
    
    add_heading(doc, 'Algorithm 2: Attendance Calculation Engine', level=2, size=14)
    add_paragraph(doc, "This algorithm calculates attendance percentages and determines classes needed or classes that can be missed based on a target goal.")

    add_heading(doc, 'Algorithm 3: Focus Session State Machine', level=2, size=14)
    add_paragraph(doc, "This manages the state transitions (STUDYING, SHORT_BREAK, LONG_BREAK) for the Pomodoro timer in Focus Mode.")

    add_heading(doc, 'Algorithm 4: AI Smart Advisor', level=2, size=14)
    add_paragraph(doc, "This flow dictates how academic context is gathered and sent to the Google Gemini API to generate personalized advice.")

    add_heading(doc, 'Algorithm 5: Study Task Lifecycle Management', level=2, size=14)
    add_paragraph(doc, "This algorithm handles task creation, automatic priority calculation based on deadlines, completion, and retrieval of overdue/pending tasks.")

    add_heading(doc, '2.2 HARDWARE AND SOFTWARE REQUIREMENTS', level=1, size=16)
    add_heading(doc, '2.2.1 HARDWARE REQUIREMENTS', level=2, size=14)
    table_hw = doc.add_table(rows=7, cols=2)
    table_hw.style = 'Table Grid'
    hw_data = [
        ("Component", "Minimum Requirement"),
        ("RAM", "4 GB RAM (8 GB recommended)"),
        ("Hard Drive", "10 GB free disk space"),
        ("Processor", "Intel Core i3 or equivalent (64-bit)"),
        ("Display", "1366 × 768 resolution or higher"),
        ("Network", "Internet connection (for MongoDB Atlas, Google OAuth, and AI API)"),
        ("Input Devices", "Keyboard and Mouse/Trackpad")
    ]
    for i, (col1, col2) in enumerate(hw_data):
        table_hw.cell(i, 0).text = col1
        table_hw.cell(i, 1).text = col2
        apply_font(table_hw.cell(i, 0).paragraphs[0].runs[0], bold=(i==0))
        apply_font(table_hw.cell(i, 1).paragraphs[0].runs[0], bold=(i==0))
        
    doc.add_paragraph()
    
    add_heading(doc, '2.2.2 SOFTWARE REQUIREMENTS', level=2, size=14)
    table_sw = doc.add_table(rows=25, cols=2)
    table_sw.style = 'Table Grid'
    sw_data = [
        ("Software", "Version / Details"),
        ("Operating System", "Windows 10/11, macOS 12+, or Linux (Ubuntu 20.04+)"),
        ("Node.js", "v18.0.0 or higher (JavaScript runtime)"),
        ("npm", "v9.0.0 or higher (Package manager)"),
        ("MongoDB Atlas", "Cloud-hosted MongoDB database (M0 free tier or higher)"),
        ("Web Browser", "Google Chrome 90+, Firefox 88+, Safari 14+, or Edge 90+"),
        ("Code Editor", "Visual Studio Code (recommended)"),
        ("Git", "v2.30+ (Version control)"),
        ("React", "v19.2.0 (Frontend framework)"),
        ("TypeScript", "v5.9.3 (Static typing for JavaScript)"),
        ("Vite", "v7.2.4 (Build tool and dev server)"),
        ("Express.js", "v4.21.0 (Backend REST API framework)"),
        ("Mongoose", "v8.7.0 (MongoDB ODM)"),
        ("Tailwind CSS", "v3.4.19 (Utility-first CSS framework)"),
        ("Framer Motion", "v12.35.0 (Animation library)"),
        ("Radix UI", "Various packages (Accessible UI primitives)"),
        ("Recharts", "v2.15.4 (Charting library)"),
        ("bcryptjs", "v2.4.3 (Password hashing)"),
        ("jsonwebtoken", "v9.0.2 (JWT authentication)"),
        ("Google Generative AI", "v0.24.1 (Gemini AI API SDK)"),
        ("google-auth-library", "v10.6.1 (Google OAuth verification)"),
        ("Nodemailer", "v8.0.1 (Email sending for OTP/reset)"),
        ("Lucide React", "v0.562.0 (Icon library)"),
        ("React Router DOM", "v7.13.1 (Client-side routing)"),
        ("Deployment", "Vercel / Netlify (Frontend), Cloud server (Backend)")
    ]
    for i, (col1, col2) in enumerate(sw_data):
        table_sw.cell(i, 0).text = col1
        table_sw.cell(i, 1).text = col2
        apply_font(table_sw.cell(i, 0).paragraphs[0].runs[0], bold=(i==0))
        apply_font(table_sw.cell(i, 1).paragraphs[0].runs[0], bold=(i==0))

    doc.add_paragraph()
    add_heading(doc, '2.3 DESIGN DETAILS', level=1, size=16)
    add_heading(doc, '2.3.1 Entity-Relationship (ER) Diagram', level=2, size=14)
    add_paragraph(doc, "The EduTrack database consists of 15 MongoDB collections (entities), handling relationships seamlessly across User, Subject, Attendance, Task, FocusSession, ChatSession, Timetable, Notifications, etc.")
    
    add_heading(doc, '2.3.2 Data Flow Diagram (DFD Level 0 – Context Diagram)', level=2, size=14)
    add_paragraph(doc, "The context diagram illustrates the student interacting with the EduTrack system, which interfaces with MongoDB Atlas, Google OAuth API, and Google Gemini AI API.")

    add_heading(doc, '2.3.3 Data Flow Diagram (DFD Level 1)', level=2, size=14)
    add_paragraph(doc, "Level 1 DFD further breaks down the student interactions into modules such as Auth, Attendance Tracker, Study Planner, Focus Mode, AI Advisor, Progress Tracker, generating queries/updates to various collections inside MongoDB Atlas.")

    doc.add_page_break()

    # Chapter 3
    add_heading(doc, 'CHAPTER 3', level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading(doc, 'IMPLEMENTATION AND RESULTS', level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_heading(doc, '3.1 IMPLEMENTATION', level=1, size=16)
    add_heading(doc, '3.1.1 Project Setup and Configuration', level=2, size=14)
    add_paragraph(doc, "The project is organized into two main directories: Frontend root directory (React + TypeScript) and server/ backend directory (Node.js + Express). It uses Vite for building and React Router DOM v7 for lazy-loaded client routing.")
    
    add_heading(doc, '3.1.2 Authentication Implementation', level=2, size=14)
    add_paragraph(doc, "Supports three methods: Email/Password (bcrypt-hashed), Google OAuth 2.0 (Google Identity Services library + Server verification), and Password Reset via Nodemailer OTP. Tokens are JWT-based.")

    add_heading(doc, '3.1.3 Backend API Architecture', level=2, size=14)
    add_paragraph(doc, "The server exposes 19 RESTful API routes modularized into auth, subjects, attendance, tasks, resources, syllabus, profile, focus, AI routes, etc., all protected by JWT middleware.")

    add_heading(doc, '3.1.4 Database Schema Design', level=2, size=14)
    add_paragraph(doc, "Mongoose schemas are used for 15 collections including User (15 fields), Attendance (Mixed type subjects and extraClasses arrays), StudyTask (16 fields including recurrence), Resource, etc.")
    
    add_paragraph(doc, "3.1.5 Frontend Component Architecture", bold=True)
    add_paragraph(doc, "The frontend component hierarchy includes ui primitives (Radix UI), Layout components, dashboard, attendance modules, gamification modules, and numerous custom hooks under the hood for API interaction and focus engine management.")
    
    add_paragraph(doc, "3.1.6 Key Custom Hooks", bold=True)
    add_paragraph(doc, "Central logic goes through custom hooks such as useData (stores synced application state), useFocusEngine (timer state machine), useSmartAcademicAssistant (Gemini AI chat), and useImportExport (portable JSON backups).")

    add_heading(doc, '3.2 RESULTS', level=1, size=16)
    result_subs = [
        ("3.2.1 Login Page", "Features a modern glassmorphism design with gradient backgrounds, Google Sign-In button, and form validations."),
        ("3.2.2 Dashboard", "Displays personalized welcome section, data cards for tasks, exam countdowns, timetable widget, and mini attendance charts."),
        ("3.2.3 Attendance Tracker", "Allows students to manage subject attendance and calculates how many classes remain to hit their target."),
        ("3.2.4 Study Planner", "Tasks, exams, and study sessions with filtering by pending/completed. Flags overdue items visually."),
        ("3.2.5 Focus Mode", "Pomodoro timer configured with ambient sounds, visually tracking progress and saving sessions to grant XP."),
        ("3.2.6 AI Smart Advisor", "Chat interface with Google Gemini offering personalized advice from application-stored context."),
        ("3.2.7 Settings", "Extensive customization: from dark/light theme options to academic data resetting and Notification Preferences.")
    ]
    for subt, desc in result_subs:
        add_heading(doc, subt, level=2, size=14)
        add_paragraph(doc, desc)

    doc.add_page_break()

    # Chapter 4
    add_heading(doc, 'CHAPTER 4', level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading(doc, 'CONCLUSION AND FUTURE SCOPE', level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_heading(doc, '4.1 CONCLUSION', level=1, size=16)
    add_paragraph(doc, "The EduTrack – Student Assistant project has been successfully developed as a comprehensive, full-stack web application. We achieved our primary objectives by creating a unified platform that integrates attendance tracking, studying planning, focus mode, and AI advice into an accessible and gamified Progressive Web App. Security is robust through multi-method authentication and password hashing. The UI is modern, using glassmorphism aesthetics to deliver a highly premium user experience.")
    
    add_heading(doc, '4.2 FUTURE SCOPE', level=1, size=16)
    add_paragraph(doc, "1. Mobile Native Application: Develop iOS/Android specific builds using React Native.")
    add_paragraph(doc, "2. Collaborative Study Groups: Allow shared progress and study resources among peer groups.")
    add_paragraph(doc, "3. Teacher/Admin Portal: Allow teachers to manage resources side by side.")
    add_paragraph(doc, "4. Advanced Analytics: Deploy ML models to predict exam performance accurately.")
    add_paragraph(doc, "5. Calendar Integration: Sync seamlessly with native calendar apps.")
    add_paragraph(doc, "6. Offline-First Architecture: Augment the service worker behavior significantly using IndexedDB.")
    add_paragraph(doc, "7. Multi-Language Support (i18n): Localize the platform for multiple regional languages.")

    doc.add_page_break()

    # Chapter 5
    add_heading(doc, 'CHAPTER 5', level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading(doc, 'REFERENCES', level=0, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    refs = [
        "1. React Documentation – Official React 19 documentation.",
        "2. TypeScript Handbook – Official TypeScript documentation.",
        "3. Node.js Documentation – Official Node.js v18+ documentation.",
        "4. Express.js Guide – Express.js web framework documentation.",
        "5. MongoDB Manual – Official MongoDB documentation.",
        "6. Mongoose ODM – Mongoose documentation for MongoDB object modeling.",
        "7. Tailwind CSS – Utility-first CSS framework documentation.",
        "8. Radix UI – Accessible component primitives documentation.",
        "9. Vite Build Tool – Next generation frontend tooling documentation.",
        "10. Google Gemini AI API – Google's Generative AI documentation.",
        "11. JSON Web Tokens (JWT) – RFC 7519 - JSON Web Token specification.",
        "12. Google OAuth 2.0 – Google Identity Services documentation.",
        "13. Framer Motion – Animation library for React.",
        "14. bcrypt.js – Password hashing library documentation.",
        "15. Recharts – React charting library documentation."
    ]
    for ref in refs:
        add_paragraph(doc, ref)
        
    doc.save('EduTrack_Project_Report.docx')
    print("Document successfully created: EduTrack_Project_Report.docx")

if __name__ == '__main__':
    main()
